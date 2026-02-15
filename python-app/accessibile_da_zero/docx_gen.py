from __future__ import annotations

import datetime as _dt
from pathlib import Path

from docx import Document


def write_accessible_docx(out_path: str | Path, titolo: str, lingua: str = "it-IT") -> Path:
    """
    Crea un DOCX con una struttura "buona" di partenza:
    - titolo
    - heading (stili Word)
    - liste
    - tabella con prima riga come intestazione (visiva)
    - istruzioni per testo alternativo immagini

    Nota: alcune proprieta' di accessibilita' (es. lang a livello XML, header row "vera") dipendono da Word/LibreOffice.
    Qui creiamo il massimo possibile senza richiedere automazioni proprietarie.
    """

    out = Path(out_path).resolve()
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Metadati base
    doc.core_properties.title = titolo
    doc.core_properties.author = "Germano Costi"
    doc.core_properties.comments = f"Template accessibile. Lingua prevista: {lingua}."

    now = _dt.datetime.now().isoformat(timespec="seconds")

    # Titolo (stile "Title")
    doc.add_paragraph(titolo, style="Title")
    doc.add_paragraph(f"Generato il {now}.", style="Subtitle")

    doc.add_heading("Come usare questo template", level=1)
    doc.add_paragraph(
        "Usa sempre i titoli (Heading 1/2/3) per la struttura. "
        "Evita di simulare titoli con testo grande o grassetto.",
        style="Normal",
    )
    doc.add_paragraph(
        "Per le immagini: aggiungi sempre una descrizione (testo alternativo). "
        "Se e' decorativa, marcalo come decorativa.",
        style="Normal",
    )
    doc.add_paragraph(
        "Per le tabelle: assicurati che le intestazioni siano impostate come riga di intestazione.",
        style="Normal",
    )

    doc.add_heading("Sezione esempio", level=1)
    doc.add_paragraph("Esempio di elenco puntato:", style="Normal")
    for item in [
        "Link descrittivi (evitare: clicca qui)",
        "Immagini con testo alternativo",
        "Tabelle con intestazioni",
        "Contrasto sufficiente (non solo colore)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Link (testo descrittivo)", level=2)
    doc.add_paragraph(
        "Scrivi link con testo significativo, ad esempio: 'Linee guida AgID' invece di 'qui'.",
        style="Normal",
    )

    doc.add_heading("Immagine (placeholder)", level=2)
    doc.add_paragraph(
        "IMMAGINE: inserisci qui l'immagine e compila il testo alternativo (alt).",
        style="Intense Quote",
    )

    doc.add_heading("Tabella (esempio)", level=2)
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Campo"
    hdr[1].text = "Valore"
    # Intestazione visiva (grassetto)
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    rows = [
        ("Versione", "0.1.0"),
        ("Lingua prevista", lingua),
        ("Nota", "Imposta la lingua in Word/LibreOffice se necessario."),
    ]
    for a, b in rows:
        row = table.add_row().cells
        row[0].text = a
        row[1].text = b

    doc.add_heading("Checklist rapida", level=1)
    for item in [
        "Titolo e metadati presenti",
        "Titoli con stili (Heading)",
        "Immagini con alt",
        "Tabelle con intestazioni",
        "Link descrittivi",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_paragraph("")
    doc.add_paragraph("Creato con Accessibile Da Zero (Germano Costi).", style="Normal")

    doc.save(str(out))
    return out

